import streamlit as st
import pandas as pd
from generate_diet import generate_diet_plan
from io import BytesIO
from docx import Document

# === Streamlit Config ===
st.set_page_config(page_title="🥗 AI Diet Planner", layout="wide")

# === Styling ===
st.markdown("""
    <style>
    .header {
        font-size: 38px;
        font-weight: 700;
        color: #2c3e50;
        margin-bottom: 20px;
    }
    .output-box {
        background-color: #ffffff;
        border: 1px solid #ddd;
        border-radius: 10px;
        padding: 20px;
        max-height: 600px;
        overflow-y: auto;
        font-family: "Courier New", monospace;
        white-space: pre-wrap;
        color: #333;
    }
    </style>
""", unsafe_allow_html=True)

st.markdown('<div class="header">🥗 AI Diet Plan Generator</div>', unsafe_allow_html=True)
st.markdown("Use AI to create a **7-day custom diet plan** based on your goals, health, and preferences.")

# === Sidebar Form ===
with st.sidebar.form("diet_form"):
    st.header("🧾 Your Details")
    age = st.number_input("Age", 10, 100, 30)
    gender = st.selectbox("Gender", ["Male", "Female", "Other"])
    height = st.number_input("Height (cm)", 100, 250, 170)
    weight = st.number_input("Weight (kg)", 30, 200, 70)
    goal = st.selectbox("Goal", ["Weight Loss", "Muscle Gain", "Eat Healthy", "Manage Condition", "General Wellness"])
    activity = st.selectbox("Activity Level", ["Sedentary", "Moderately Active", "Active"])
    food_pref = st.selectbox("Diet Preference", ["Vegetarian", "Non-Vegetarian", "Vegan"])

    cuisine = st.selectbox("Preferred Cuisine Type", ["None", "Indian", "American", "Mediterranean", "Asian", "Global"])
    cuisine_selected = None if cuisine == "None" else cuisine

    conditions = st.text_area("Health Conditions", placeholder="e.g., diabetes, hypertension, thyroid...")
    allergies = st.text_input("Allergies (if any)", placeholder="e.g., nuts, gluten, lactose...")
    meals = st.slider("Preferred Meals per Day", 1, 6, 3)

    if goal == "Weight Loss":
        weight_to_lose = st.number_input("Target Weight Loss (kg)", 1, 50, 5)
        duration_days = st.number_input("Timeframe to Achieve Goal (days)", 7, 180, 30)
    else:
        weight_to_lose = 0
        duration_days = 30

    submitted = st.form_submit_button("Generate Plan")

# === Function to parse plan text to DataFrame ===
def parse_diet_to_df(plan_text: str) -> pd.DataFrame:
    days = plan_text.strip().split("Day ")[1:]  # splits into day blocks
    rows = []
    for day_text in days:
        lines = day_text.strip().splitlines()
        day_num = lines[0].strip(":").strip()  # e.g. "1"
        meal_data = {}
        for line in lines[1:]:
            if line.startswith("Meal"):
                if meal_data:
                    rows.append(meal_data)
                    meal_data = {}
                meal_data = {"Day": f"Day {day_num}", "Meal": line.strip()}
            elif "Dish:" in line:
                meal_data["Dish"] = line.split("Dish:")[1].strip()
            elif "Time:" in line:
                meal_data["Time"] = line.split("Time:")[1].strip()
            elif "Calories:" in line:
                meal_data["Calories"] = line.split("Calories:")[1].strip()
            elif "Purpose:" in line or "Reason:" in line:
                if "Purpose:" in line:
                    meal_data["Purpose"] = line.split("Purpose:")[1].strip()
                else:
                    meal_data["Purpose"] = line.split("Reason:")[1].strip()
        if meal_data:
            rows.append(meal_data)
    return pd.DataFrame(rows)

# === Function to generate Word doc from DataFrame ===
def generate_word_doc(df: pd.DataFrame) -> BytesIO:
    doc = Document()
    doc.add_heading("7-Day Personalized Diet Plan", level=1)
    for day in df["Day"].unique():
        doc.add_heading(day, level=2)
        sub_df = df[df["Day"] == day]
        for _, row in sub_df.iterrows():
            doc.add_paragraph(f'{row["Meal"]}')
            doc.add_paragraph(f'Dish: {row["Dish"]}')
            doc.add_paragraph(f'Time: {row["Time"]}')
            doc.add_paragraph(f'Calories: {row["Calories"]} kcal')
            doc.add_paragraph(f'Purpose: {row["Purpose"]}')
            doc.add_paragraph("")
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# === Main logic after form submission ===
if submitted:
    # Calculate TDEE and target calories
    if activity == "Sedentary":
        multiplier = 1.2
    elif activity == "Moderately Active":
        multiplier = 1.55
    else:
        multiplier = 1.725

    bmr = 10 * weight + 6.25 * height - 5 * age + (5 if gender == "Male" else -161)
    tdee = bmr * multiplier

    total_deficit = weight_to_lose * 7700
    daily_deficit = total_deficit / duration_days if duration_days else 0
    target_daily_calories = int(tdee - daily_deficit) if goal == "Weight Loss" else int(tdee)

    user_input = f"""
    {gender}, {age} years old, {height} cm, {weight} kg.
    Goal: {goal}, Diet: {food_pref}, Activity: {activity}.
    Preferred meals per day: {meals}.
    Health conditions: {conditions or 'None'}.
    Allergies: {allergies or 'None'}.
    Suggested Daily Calorie Limit: {target_daily_calories} kcal
    """

    with st.spinner("Generating your personalized 7-day diet plan... Please wait ⏳"):
        plan = generate_diet_plan(user_input, cuisine=cuisine_selected)

    st.success("✅ Your personalized 7-day diet plan:")
    st.markdown(f'<div class="output-box">{plan}</div>', unsafe_allow_html=True)

    # Parse plan text to DataFrame
    df = parse_diet_to_df(plan)
    df["Calories"] = pd.to_numeric(df["Calories"], errors="coerce")
    calorie_summary = df.groupby("Day")["Calories"].sum().reset_index()

    st.subheader("📊 Daily Calorie Summary")
    st.dataframe(calorie_summary)

    # Generate Excel file
    excel_output = BytesIO()
    with pd.ExcelWriter(excel_output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Diet Plan')
        calorie_summary.to_excel(writer, index=False, sheet_name='Calories Summary')
    excel_output.seek(0)
    st.session_state.excel_output = excel_output.getvalue()

    # Generate Word file
    docx_output = generate_word_doc(df)
    st.session_state.docx_output = docx_output.getvalue()

# === Show download buttons only if outputs exist ===
if st.session_state.get("excel_output") and st.session_state.get("docx_output"):
    st.download_button(
        "📥 Download as Excel",
        data=st.session_state.excel_output,
        file_name="diet_plan.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        key="download_excel"
    )

    st.download_button(
        "📄 Download as Word (.docx)",
        data=st.session_state.docx_output,
        file_name="diet_plan.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="download_word"
    )

# === Clear outputs button ===
if st.button("Clear Output"):
    st.session_state.excel_output = None
    st.session_state.docx_output = None
    st.experimental_rerun()
